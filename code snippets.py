from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("APPENDIX A\n\nCODE SNIPPETS")
run.font.size = Pt(16)
run.font.bold = True

doc.add_paragraph()

# A.1 Data Loading
doc.add_heading("A.1  Data Loading", level=2)

code1 = '''import pandas as pd
import numpy as np

def load_german_credit_data(filepath='data/raw/credit_data.csv'):
    """
    Load German Credit Dataset
    
    Returns:
        X: Feature matrix (numpy array)
        y: Target vector (numpy array)
    """
    # Load data
    df = pd.read_csv(filepath)
    
    # Separate features and target
    X = df.drop('credit_risk', axis=1).values
    y = df['credit_risk'].values
    
    return X, y
'''

p = doc.add_paragraph(code1)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.2 Data Preprocessing
doc.add_heading("A.2  Data Preprocessing with SMOTE", level=2)

code2 = '''from sklearn.preprocessing import StandardScaler
from imblearn.over_sampling import SMOTE

def preprocess_data(X_train, y_train, use_smote=True):
    """
    Preprocess training data with scaling and SMOTE
    
    Args:
        X_train: Training features
        y_train: Training labels
        use_smote: Whether to apply SMOTE
    
    Returns:
        X_train_processed: Preprocessed features
        y_train_processed: Processed labels
    """
    # Standard scaling
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    
    # Apply SMOTE if requested
    if use_smote:
        smote = SMOTE(random_state=42, k_neighbors=5)
        X_train_processed, y_train_processed = smote.fit_resample(
            X_train_scaled, y_train
        )
        print(f"After SMOTE: {len(y_train_processed)} samples")
    else:
        X_train_processed = X_train_scaled
        y_train_processed = y_train
    
    return X_train_processed, y_train_processed, scaler
'''

p = doc.add_paragraph(code2)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.3 Random Forest Training
doc.add_heading("A.3  Random Forest Training", level=2)

code3 = '''from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report

def train_random_forest(X_train, y_train, X_test, y_test):
    """
    Train Random Forest classifier
    
    Returns:
        model: Trained Random Forest model
        metrics: Performance metrics dictionary
    """
    # Initialize model
    rf_model = RandomForestClassifier(
        n_estimators=100,
        max_depth=20,
        min_samples_split=5,
        min_samples_leaf=2,
        class_weight='balanced',
        random_state=42
    )
    
    # Train model
    rf_model.fit(X_train, y_train)
    
    # Predictions
    y_pred = rf_model.predict(X_test)
    
    # Calculate metrics
    accuracy = accuracy_score(y_test, y_pred)
    
    print(f"Random Forest Accuracy: {accuracy*100:.2f}%")
    print(classification_report(y_test, y_pred))
    
    return rf_model, accuracy
'''

p = doc.add_paragraph(code3)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.4 Neural Network Architecture
doc.add_heading("A.4  Neural Network Architecture", level=2)

code4 = '''import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers

def create_neural_network(input_dim):
    """
    Create neural network for credit risk prediction
    
    Args:
        input_dim: Number of input features
    
    Returns:
        model: Compiled Keras model
    """
    model = keras.Sequential([
        layers.Dense(64, activation='relu', input_shape=(input_dim,)),
        layers.Dropout(0.3),
        layers.Dense(32, activation='relu'),
        layers.Dropout(0.3),
        layers.Dense(1, activation='sigmoid')
    ])
    
    model.compile(
        optimizer='adam',
        loss='binary_crossentropy',
        metrics=['accuracy']
    )
    
    return model

# Training
model = create_neural_network(input_dim=24)
history = model.fit(
    X_train, y_train,
    validation_split=0.2,
    epochs=50,
    batch_size=32,
    verbose=1
)
'''

p = doc.add_paragraph(code4)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.5 Federated Client
doc.add_heading("A.5  Federated Learning Client", level=2)

code5 = '''class FederatedClient:
    """Client for federated learning"""
    
    def __init__(self, client_id, X_train, y_train, X_test, y_test):
        self.client_id = client_id
        self.X_train = X_train
        self.y_train = y_train
        self.X_test = X_test
        self.y_test = y_test
        self.model = None
    
    def build_model(self, input_dim):
        """Build local neural network model"""
        self.model = keras.Sequential([
            layers.Dense(64, activation='relu', input_shape=(input_dim,)),
            layers.Dropout(0.3),
            layers.Dense(32, activation='relu'),
            layers.Dropout(0.3),
            layers.Dense(1, activation='sigmoid')
        ])
        self.model.compile(
            optimizer='adam',
            loss='binary_crossentropy',
            metrics=['accuracy']
        )
    
    def train(self, epochs=5):
        """Train model on local data"""
        history = self.model.fit(
            self.X_train, self.y_train,
            epochs=epochs,
            batch_size=32,
            verbose=0
        )
        return history
    
    def get_weights(self):
        """Get current model weights"""
        return self.model.get_weights()
    
    def set_weights(self, weights):
        """Update model with new weights"""
        self.model.set_weights(weights)
    
    def evaluate(self):
        """Evaluate model on local test data"""
        loss, accuracy = self.model.evaluate(
            self.X_test, self.y_test, verbose=0
        )
        return accuracy
'''

p = doc.add_paragraph(code5)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.6 FedAvg Algorithm
doc.add_heading("A.6  FedAvg Aggregation", level=2)

code6 = '''def federated_averaging(client_weights, client_sizes):
    """
    Aggregate client weights using FedAvg algorithm
    
    Args:
        client_weights: List of weight arrays from each client
        client_sizes: List of dataset sizes for each client
    
    Returns:
        averaged_weights: Aggregated model weights
    """
    total_size = sum(client_sizes)
    
    # Initialize averaged weights with zeros
    averaged_weights = [np.zeros_like(w) for w in client_weights[0]]
    
    # Weighted average
    for client_w, size in zip(client_weights, client_sizes):
        weight = size / total_size
        for i, w in enumerate(client_w):
            averaged_weights[i] += w * weight
    
    return averaged_weights

# Example usage
global_weights = federated_averaging(
    client_weights=[client1.get_weights(), 
                   client2.get_weights(),
                   client3.get_weights()],
    client_sizes=[len(client1.X_train),
                 len(client2.X_train),
                 len(client3.X_train)]
)
'''

p = doc.add_paragraph(code6)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

# A.7 Evaluation Metrics
doc.add_heading("A.7  Performance Evaluation", level=2)

code7 = '''from sklearn.metrics import accuracy_score, precision_score, recall_score
from sklearn.metrics import f1_score, roc_auc_score, confusion_matrix

def calculate_metrics(y_true, y_pred, y_pred_proba):
    """
    Calculate comprehensive performance metrics
    
    Returns:
        metrics: Dictionary of evaluation metrics
    """
    metrics = {
        'accuracy': accuracy_score(y_true, y_pred),
        'precision': precision_score(y_true, y_pred),
        'recall': recall_score(y_true, y_pred),
        'f1_score': f1_score(y_true, y_pred),
        'auc_roc': roc_auc_score(y_true, y_pred_proba),
        'confusion_matrix': confusion_matrix(y_true, y_pred)
    }
    
    return metrics

# Usage
y_pred = model.predict(X_test) > 0.5
y_pred_proba = model.predict(X_test)
metrics = calculate_metrics(y_test, y_pred, y_pred_proba)

print(f"Accuracy:  {metrics['accuracy']*100:.2f}%")
print(f"Precision: {metrics['precision']:.4f}")
print(f"Recall:    {metrics['recall']:.4f}")
print(f"F1-Score:  {metrics['f1_score']:.4f}")
print(f"AUC-ROC:   {metrics['auc_roc']:.4f}")
'''

p = doc.add_paragraph(code7)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(9)

doc.save('appendix_a_code_snippets.docx')
print("✅ Appendix A created!")